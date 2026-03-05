from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_word_document(paper_data: dict) -> bytes:
    """根据文稿数据生成标准格式Word文档"""
    
    doc = Document()
    
    set_document_styles(doc)
    
    title = paper_data.get('title', '无标题文稿')
    keywords = paper_data.get('keywords', [])
    
    add_cover_page(doc, title, paper_data)
    
    add_abstract_chinese(doc, paper_data)
    
    add_abstract_english(doc, title, keywords)
    
    add_table_of_contents(doc, paper_data)
    
    add_chapters(doc, paper_data)
    
    add_references(doc, paper_data)
    
    add_appendix(doc, paper_data)
    
    return doc

def add_cover_page(doc: Document, title: str, paper_data: dict):
    """添加封面页"""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    info_items = [
        ("院    系", paper_data.get('major', 'XXX学院')),
        ("专    业", paper_data.get('major', 'XXX专业')),
        ("学生姓名", paper_data.get('author', 'XXX')),
        ("学    号", paper_data.get('student_id', 'XXXXXXXXXX')),
        ("指导教师", paper_data.get('advisor', 'XXX')),
    ]
    
    for label, value in info_items:
        para = doc.add_paragraph()
        para.add_run(f"  {label}：{value}")
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()

def add_abstract_chinese(doc: Document, paper_data: dict):
    """添加中文摘要"""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("摘    要")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(18)
    
    abstract_content = paper_data.get('abstract_cn', '')
    if not abstract_content:
        abstract_content = generate_abstract_from_content(paper_data, 'chinese')
    
    if abstract_content:
        abstract_para = doc.add_paragraph()
        abstract_run = abstract_para.add_run(abstract_content)
        abstract_run.font.size = Pt(12)
        abstract_para.paragraph_format.first_line_indent = Inches(0.3)
        abstract_para.paragraph_format.line_spacing = 1.5
        abstract_para.paragraph_format.space_after = Pt(12)
    
    keywords_text = '，'.join(paper_data.get('keywords', []))
    if keywords_text:
        kw_para = doc.add_paragraph()
        kw_run = kw_para.add_run(f"关键词：{keywords_text}")
        kw_run.font.size = Pt(12)
        kw_run.bold = True

def add_abstract_english(doc: Document, title: str, keywords: list):
    """添加英文摘要"""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("ABSTRACT")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(18)
    
    eng_title_para = doc.add_paragraph()
    eng_title_run = eng_title_para.add_run(title)
    eng_title_run.font.size = Pt(14)
    eng_title_run.bold = True
    eng_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eng_title_para.paragraph_format.space_after = Pt(12)
    
    abstract_para = doc.add_paragraph()
    abstract_run = abstract_para.add_run("This paper...")
    abstract_run.font.size = Pt(12)
    abstract_para.paragraph_format.first_line_indent = Inches(0.3)
    abstract_para.paragraph_format.line_spacing = 1.5
    abstract_para.paragraph_format.space_after = Pt(12)
    
    if keywords:
        kw_para = doc.add_paragraph()
        kw_run = kw_para.add_run(f"Keywords: {', '.join(keywords)}")
        kw_run.font.size = Pt(12)
        kw_run.bold = True

def add_table_of_contents(doc: Document, paper_data: dict):
    """添加目录"""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("目    录")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(18)
    
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.line_spacing = 2
    
    outline = paper_data.get('outline', [])
    chapter_num = 1
    
    for chapter in outline:
        chapter_title = chapter.get('title', '')
        if not chapter_title or chapter_title == '参考文献':
            continue
            
        if chapter_title != '附录':
            dot_line = " .................................................... "
            run = toc_para.add_run(f"第{chapter_num}章    {chapter_title}{dot_line}{chapter_num}")
            run.font.size = Pt(12)
            chapter_num += 1
            
            for section in chapter.get('sections', []):
                section_title = section.get('title', '')
                if section_title:
                    sec_run = toc_para.add_run(f"    {section_title}\n")
                    sec_run.font.size = Pt(12)
    
    ref_run = toc_para.add_run(f"参考文献{dot_line}{chapter_num}")
    ref_run.font.size = Pt(12)
    
    doc.add_page_break()

def add_chapters(doc: Document, paper_data: dict):
    """添加章节内容"""
    outline = paper_data.get('outline', [])
    chapter_num = 1
    
    for chapter in outline:
        chapter_title = chapter.get('title', '')
        if not chapter_title or chapter_title == '参考文献' or chapter_title == '附录':
            continue
            
        ch_para = doc.add_paragraph()
        ch_run = ch_para.add_run(f"第{chapter_num}章    {chapter_title}")
        ch_run.font.size = Pt(16)
        ch_run.bold = True
        ch_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ch_para.paragraph_format.space_before = Pt(24)
        ch_para.paragraph_format.space_after = Pt(18)
        
        for section in chapter.get('sections', []):
            section_title = section.get('title', '')
            section_content = section.get('content', '')
            
            if section_title:
                sec_para = doc.add_paragraph()
                sec_run = sec_para.add_run(section_title)
                sec_run.font.size = Pt(14)
                sec_run.bold = True
                sec_para.paragraph_format.space_before = Pt(12)
                sec_para.paragraph_format.space_after = Pt(6)
                
                if section_content:
                    content_para = doc.add_paragraph()
                    formatted_content = format_content_for_word(section_content)
                    content_run = content_para.add_run(formatted_content)
                    content_run.font.size = Pt(12)
                    content_para.paragraph_format.first_line_indent = Inches(0.3)
                    content_para.paragraph_format.line_spacing = 1.5
                    content_para.paragraph_format.space_after = Pt(12)
        
        chapter_num += 1

def add_references(doc: Document, paper_data: dict):
    """添加参考文献"""
    doc.add_page_break()
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("参考文献")
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(18)
    
    references = paper_data.get('references', [])
    if not references:
        references = get_default_references()
    
    for i, ref in enumerate(references, 1):
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run(f"[{i}] {ref}")
        ref_run.font.size = Pt(11)
        ref_para.paragraph_format.first_line_indent = Inches(-0.3)
        ref_para.paragraph_format.left_indent = Inches(0.3)
        ref_para.paragraph_format.line_spacing = 1.5
        ref_para.paragraph_format.space_after = Pt(6)

def add_appendix(doc: Document, paper_data: dict):
    """添加附录"""
    outline = paper_data.get('outline', [])
    
    has_appendix = any(ch.get('title') == '附录' for ch in outline)
    
    if has_appendix:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("附    录")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(24)
        title_para.paragraph_format.space_after = Pt(18)
        
        for chapter in outline:
            if chapter.get('title') == '附录':
                for section in chapter.get('sections', []):
                    section_title = section.get('title', '')
                    section_content = section.get('content', '')
                    
                    if section_title:
                        sec_para = doc.add_paragraph()
                        sec_run = sec_para.add_run(section_title)
                        sec_run.font.size = Pt(14)
                        sec_run.bold = True
                        
                        if section_content:
                            content_para = doc.add_paragraph()
                            content_run = content_para.add_run(section_content)
                            content_run.font.size = Pt(12)
                            content_para.paragraph_format.first_line_indent = Inches(0.3)
                            content_para.paragraph_format.line_spacing = 1.5

def generate_abstract_from_content(paper_data: dict, lang: str = 'chinese') -> str:
    """从内容生成摘要"""
    outline = paper_data.get('outline', [])
    
    all_content = []
    for chapter in outline:
        for section in chapter.get('sections', []):
            content = section.get('content', '')
            if content and len(content) > 50:
                all_content.append(content)
    
    if not all_content:
        return ""
    
    combined = ' '.join(all_content[:3])
    
    if lang == 'chinese':
        abstract = combined[:500] + "..." if len(combined) > 500 else combined
        return abstract
    else:
        abstract = "This paper studies " + combined[:400] + "..." if len(combined) > 400 else "This paper studies " + combined
        return abstract

def get_default_references() -> list:
    """获取默认参考文献格式"""
    return [
        "[1] 张三, 李四. 研究方法论[M]. 北京: 高等教育出版社, 2020.",
        "[2] Wang, J. Research on Academic Writing[J]. Journal of Education, 2019, 45(2): 123-135.",
        "[3] 王五. 文稿写作规范指南[M]. 上海: 上海交通大学出版社, 2018.",
        "[4] Smith, A. & Brown, B. Academic Research Methods[M]. New York: Academic Press, 2021.",
        "[5] 赵六. 创意文稿引用与参考文献管理[J]. 图书情报工作, 2017, 61(15): 78-85.",
    ]

def set_document_styles(doc: Document):
    """设置文档样式"""
    styles = doc.styles
    
    if 'Normal' in styles:
        style = styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        para = style.paragraph_format
        para.line_spacing = 1.5
        para.first_line_indent = Inches(0.3)

def format_content_for_word(content: str) -> str:
    """格式化内容以适应Word"""
    content = content.strip()
    content = re.sub(r'\n{3,}', '\n\n', content)
    content = re.sub(r'^\s*[-•]\s*', '　　', content, flags=re.MULTILINE)
    return content

def format_paragraph_spacing(paragraph):
    """格式化段落间距"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

def export_paper_to_file(paper_data: dict, file_path: str):
    """导出文稿到指定文件路径"""
    doc = create_word_document(paper_data)
    doc.save(file_path)
